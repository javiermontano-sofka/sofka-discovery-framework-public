#!/usr/bin/env python3
"""DOCX → priming-rag. Headings + paragraphs + tables."""
from pathlib import Path
from docx import Document
from _common import cli, file_meta, write_priming, truncate


def main() -> None:
    args = cli()
    path = Path(args.path)
    doc = Document(path)
    meta = file_meta(path)

    lines: list[str] = []
    headings: list[str] = []
    for p in doc.paragraphs:
        style = (p.style.name or "").lower()
        text = p.text.strip()
        if not text:
            continue
        if style.startswith("heading"):
            level = "".join(c for c in style if c.isdigit()) or "1"
            lines.append(f"{'#' * (int(level) + 1)} {text}")
            headings.append(text)
        else:
            lines.append(text)

    for i, tbl in enumerate(doc.tables[:30], 1):
        lines.append(f"\n### Tabla {i}\n")
        for row in tbl.rows[:50]:
            cells = " | ".join(c.text.strip().replace("\n", " ") for c in row.cells)
            lines.append(f"| {cells} |")

    resumen = [
        f"DOCX con {len(doc.paragraphs)} párrafos, {len(doc.tables)} tablas, {len(headings)} headings",
        f"Headings top: {', '.join(headings[:6])}" if headings else "Sin headings",
    ]
    contenido = "\n\n".join(lines)
    evidencia = [f"[ADJUNTO:{path.name}:heading={h}]" for h in headings[:8]]
    out = write_priming(Path(args.out) if args.out else None, meta, "docx", resumen, truncate(contenido), evidencia)
    print(str(out))


if __name__ == "__main__":
    main()
